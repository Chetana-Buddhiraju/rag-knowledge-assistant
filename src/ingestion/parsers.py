"""Parse raw source documents (PDF / DOCX / XLSX) into a flat list of
`Block` objects: a paragraph/row of text tagged with the section heading it
falls under. Keeping the heading attached to every block is what lets the
chunker build section-aware chunks instead of blind fixed-size windows.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

import openpyxl
from docx import Document as DocxDocument
from pypdf import PdfReader

# A heading is a short line, title-cased or numbered ("3. Expense Categories & Limits"),
# with no trailing punctuation typical of prose sentences.
_HEADING_RE = re.compile(r"^(\d+(\.\d+)*\.?\s+)?[A-Z][A-Za-z0-9 &/(),'\-—–]{2,80}$")

# Repeated page chrome (company footer/header line, bare "Page N" markers) that
# carries no document content and must never become a section heading or a
# retrievable block — left in, it pollutes the index with near-duplicate
# "Header"/"Page N" chunks that outrank real content on short queries.
_FOOTER_RE = re.compile(r"northwind traders,?\s+inc\.?\s*[—-]?\s*internal use only", re.I)
_PAGE_MARKER_RE = re.compile(r"^page\s+\d+$", re.I)


@dataclass
class Block:
    text: str
    section: str


def _looks_like_heading(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 90:
        return False
    if line.endswith((".", ":", ";")) and not re.match(r"^\d+(\.\d+)*\.?\s", line):
        return False
    return bool(_HEADING_RE.match(line))


def parse_pdf(path: Path) -> list[Block]:
    reader = PdfReader(str(path))
    blocks: list[Block] = []
    section = "Header"
    for page in reader.pages:
        text = page.extract_text() or ""
        for raw_line in text.split("\n"):
            line = raw_line.strip()
            if not line:
                continue
            # Drop repeated page chrome before it can be mistaken for a heading
            # (a bare "Page 2" line otherwise matches the heading pattern).
            if _FOOTER_RE.search(line) or _PAGE_MARKER_RE.match(line):
                continue
            if _looks_like_heading(line):
                section = line
                continue
            blocks.append(Block(text=line, section=section))
    return _merge_consecutive(blocks)


def parse_docx(path: Path) -> list[Block]:
    doc = DocxDocument(str(path))
    blocks: list[Block] = []
    section = "Header"
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or _looks_like_heading(line):
            section = line
            continue
        blocks.append(Block(text=line, section=section))
    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        table_text = _table_to_text(rows)
        if table_text:
            blocks.append(Block(text=table_text, section=f"{section} (table)"))
    return _merge_consecutive(blocks)


def parse_xlsx(path: Path) -> list[Block]:
    wb = openpyxl.load_workbook(str(path), data_only=True)
    blocks: list[Block] = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                rows.append(["" if c is None else str(c) for c in row])
        if not rows:
            continue
        table_text = _table_to_text(rows)
        blocks.append(Block(text=table_text, section=f"Sheet: {ws.title}"))
    return blocks


def _table_to_text(rows: list[list[str]]) -> str:
    lines = [" | ".join(cell for cell in row if cell) for row in rows if any(row)]
    return "\n".join(lines)


def _merge_consecutive(blocks: list[Block]) -> list[Block]:
    """Merge adjacent lines under the same heading into one paragraph block."""
    merged: list[Block] = []
    for b in blocks:
        if merged and merged[-1].section == b.section and not merged[-1].text.endswith("(table)"):
            merged[-1] = Block(text=merged[-1].text + "\n" + b.text, section=b.section)
        else:
            merged.append(b)
    return merged


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
}


def parse_document(path: Path) -> list[Block]:
    parser = PARSERS.get(path.suffix.lower())
    if parser is None:
        raise ValueError(f"Unsupported file type: {path.suffix}")
    return parser(path)
